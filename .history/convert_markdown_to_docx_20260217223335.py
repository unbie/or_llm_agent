"""
学术论文Markdown转Word转换脚本
功能: 将Markdown文件转换为符合学术规范的Word文档
依赖: pip install python-docx markdown
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
import re


def setup_academic_styles(doc):
    """设置学术论文样式"""
    
    # 设置默认字体为中文宋体、英文Times New Roman
    doc.styles['Normal'].font.name = '宋体'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:ascii'), 'Times New Roman')
    doc.styles['Normal'].font.size = Pt(12)  # 小四
    
    # 设置段落格式
    paragraph_format = doc.styles['Normal'].paragraph_format
    paragraph_format.line_spacing = 1.5  # 1.5倍行距
    paragraph_format.space_before = Pt(0)
    paragraph_format.space_after = Pt(0)
    
    # 标题1样式（章）
    heading1 = doc.styles['Heading 1']
    heading1.font.name = '黑体'
    heading1.font.size = Pt(16)  # 三号
    heading1.font.bold = True
    heading1._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    heading1.paragraph_format.space_before = Pt(6)
    heading1.paragraph_format.space_after = Pt(6)
    
    # 标题2样式（节）
    heading2 = doc.styles['Heading 2']
    heading2.font.name = '黑体'
    heading2.font.size = Pt(15)  # 小三
    heading2.font.bold = True
    heading2._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    heading2.paragraph_format.space_before = Pt(6)
    heading2.paragraph_format.space_after = Pt(6)
    
    # 标题3样式（小节）
    heading3 = doc.styles['Heading 3']
    heading3.font.name = '黑体'
    heading3.font.size = Pt(14)  # 四号
    heading3.font.bold = True
    heading3._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    heading3.paragraph_format.space_before = Pt(6)
    heading3.paragraph_format.space_after = Pt(6)
    
    # 标题4样式
    heading4 = doc.styles['Heading 4']
    heading4.font.name = '黑体'
    heading4.font.size = Pt(12)  # 小四
    heading4.font.bold = True
    heading4._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')


def parse_markdown_file(filepath):
    """解析Markdown文件"""
    with open(filepath, 'r', encoding='utf-8') as f:
        content = f.read()
    return content


def convert_md_to_docx(md_filepath, output_filepath):
    """
    将Markdown文件转换为Word文档
    
    Args:
        md_filepath: Markdown文件路径
        output_filepath: 输出的Word文档路径
    """
    
    # 创建Word文档
    doc = Document()
    
    # 设置页面边距（上下2.54cm，左右3.17cm）
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)
    
    # 应用学术样式
    setup_academic_styles(doc)
    
    # 读取Markdown内容
    content = parse_markdown_file(md_filepath)
    lines = content.split('\n')
    
    print(f"正在转换 {md_filepath}...")
    print(f"总行数: {len(lines)}")
    
    in_code_block = False
    code_buffer = []
    in_table = False
    table_buffer = []
    
    for i, line in enumerate(lines):
        if i % 100 == 0:
            print(f"处理进度: {i}/{len(lines)} ({i*100//len(lines)}%)")
        
        # 处理代码块
        if line.strip().startswith('```'):
            if not in_code_block:
                in_code_block = True
                code_buffer = []
            else:
                # 代码块结束，添加到文档
                code_text = '\n'.join(code_buffer)
                p = doc.add_paragraph(code_text)
                p.style = 'No Spacing'
                p_format = p.paragraph_format
                p_format.left_indent = Cm(1)
                # 设置等宽字体
                for run in p.runs:
                    run.font.name = 'Courier New'
                    run.font.size = Pt(10)
                in_code_block = False
                code_buffer = []
            continue
        
        if in_code_block:
            code_buffer.append(line)
            continue
        
        # 处理标题
        heading_match = re.match(r'^(#{1,4})\s+(.+)$', line)
        if heading_match:
            level = len(heading_match.group(1))
            title = heading_match.group(2)
            doc.add_heading(title, level=level)
            continue
        
        # 处理水平线
        if line.strip() in ['---', '***', '___']:
            doc.add_paragraph('_' * 50)
            continue
        
        # 处理引用块
        if line.strip().startswith('>'):
            quote_text = line.strip()[1:].strip()
            p = doc.add_paragraph(quote_text)
            p.style = 'Intense Quote'
            continue
        
        # 处理列表（简化版本）
        if re.match(r'^\s*[-*+]\s+', line):
            list_text = re.sub(r'^\s*[-*+]\s+', '', line)
            doc.add_paragraph(list_text, style='List Bullet')
            continue
        
        if re.match(r'^\s*\d+\.\s+', line):
            list_text = re.sub(r'^\s*\d+\.\s+', '', line)
            doc.add_paragraph(list_text, style='List Number')
            continue
        
        # 处理表格（简化版本 - 建议使用Pandoc处理复杂表格）
        if '|' in line and line.strip().startswith('|'):
            # 这里只做简单标记，复杂表格建议用Pandoc
            doc.add_paragraph(line)
            continue
        
        # 处理普通段落
        if line.strip():
            # 处理粗体和斜体
            text = line
            text = re.sub(r'\*\*(.+?)\*\*', r'<b>\1</b>', text)
            text = re.sub(r'\*(.+?)\*', r'<i>\1</i>', text)
            text = re.sub(r'`(.+?)`', r'<code>\1</code>', text)
            
            p = doc.add_paragraph()
            
            # 简化处理：直接添加文本（完整版需要解析HTML标签）
            p.add_run(line)
    
    # 保存文档
    doc.save(output_filepath)
    print(f"\n转换完成！")
    print(f"输出文件: {output_filepath}")


def add_table_of_contents(doc):
    """添加目录（需要在Word中手动刷新）"""
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')
    
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
    
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    
    r_element = run._element
    r_element.append(fldChar)
    r_element.append(instrText)
    r_element.append(fldChar2)
    r_element.append(fldChar3)
    
    return paragraph


if __name__ == '__main__':
    # 转换配置
    input_file = 'ALNS_VRP_项目报告.md'
    output_file = 'ALNS_VRP_项目报告_学术版.docx'
    
    print("=" * 60)
    print("学术论文Markdown转Word工具")
    print("=" * 60)
    print()
    
    try:
        convert_md_to_docx(input_file, output_file)
        print()
        print("=" * 60)
        print("转换成功！")
        print("=" * 60)
        print()
        print("⚠️  重要提示:")
        print("1. 请在Word中打开文档，手动添加目录（引用→目录→自动目录）")
        print("2. 检查表格格式并手动调整")
        print("3. 复杂的数学公式可能需要手动调整")
        print("4. 建议使用Pandoc工具以获得更好的转换效果")
        
    except Exception as e:
        print(f"转换失败: {e}")
        print()
        print("建议使用Pandoc工具进行转换（参见 转换说明.md）")
